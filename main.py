import os
import re
import shutil

import win32com.client as win32

from tkinter.ttk import Label, OptionMenu
import tkinter as tk
from tkinter import filedialog, messagebox, StringVar, OptionMenu, Label, Button, font, simpledialog
from tkinter.font import families

from tkinter.simpledialog import askinteger
import win32com.client
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt, Mm, Cm, Inches
from docx.oxml.ns import nsmap
from docx.enum.text import WD_LINE_SPACING


def clear_com_cache():
    cache_path = os.path.join(os.environ['LOCALAPPDATA'], 'Temp', 'gen_py')
    if os.path.exists(cache_path):
        shutil.rmtree(cache_path)
    win32com.client.gencache.is_readonly = False
    win32com.client.gencache.Rebuild()


def select_document():
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(title="选择要处理的文档", filetypes=[("Word Documents", "*.doc;*.docx")])
    root.destroy()

    if file_path:
        # 检查文件扩展名
        if file_path.endswith('.doc'):
            # 转换为docx
            return convert_doc_to_docx(file_path)
        elif file_path.endswith('.docx'):
            return file_path

    return None


def select_save_as():
    root = tk.Tk()
    root.withdraw()
    # 允许用户选择保存为.doc或.docx
    file_path = filedialog.asksaveasfilename(title="另存为...", filetypes=[("Word Documents", "*.doc;*.docx")],
                                             defaultextension=".docx")
    root.destroy()
    return file_path


def convert_doc_to_docx(input_path):
    output_path = input_path + "x"  # 假设自动将文件名后添加x来转换格式
    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=16)  # FileFormat=16 for .docx
    doc.Close()
    word.Quit()
    return output_path


def has_graphics(paragraph):
    """检查段落是否包含图形（如图片、图表、公式、流程图等）。"""
    for run in paragraph.runs:
        # 直接检查XML中的关键元素
        xml_str = run._element.xml
        if '<m:oMath' in xml_str or \
                '<wp:anchor' in xml_str or '<v:shape' in xml_str:
            # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return True
        if '<pic:pic' in xml_str or '<wp:inline' in xml_str:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return True
    return False


key1 = r'^摘\s*要$'
key2 = r'^目\s*录$'


# 默认模板
def apply_default_template():
    def set_page_layout(doc):
        for section in doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.orientation = WD_ORIENT.PORTRAIT
            section.top_margin = Mm(30)
            section.bottom_margin = Mm(20)
            section.left_margin = Mm(30)
            section.right_margin = Mm(20)
            section.gutter = Mm(10)
            section.footer_distance = Cm(1)
            section.header_distance = Cm(2)

    def update_headers_if_text_exists(doc, header_text):
        for section in doc.sections:
            if any(paragraph.text.strip() for paragraph in section.header.paragraphs):
                clear_and_set_new_header(section.header, header_text)
            if not section.first_page_header.is_linked_to_previous:
                if any(paragraph.text.strip() for paragraph in section.first_page_header.paragraphs):
                    clear_and_set_new_header(section.first_page_header, header_text)
            if section.even_page_header and not section.even_page_header.is_linked_to_previous:
                if any(paragraph.text.strip() for paragraph in section.even_page_header.paragraphs):
                    clear_and_set_new_header(section.even_page_header, header_text)

    def clear_and_set_new_header(header, text):
        for paragraph in header.paragraphs:
            paragraph.clear()
        new_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        new_paragraph.different_first_page_header_footer = False
        run = new_paragraph.add_run(text)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def format_abstract(doc, text_to_format):
        for paragraph in doc.paragraphs:
            if text_to_format in paragraph.text:
                for run in paragraph.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(16)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_before = Pt(28)
                paragraph.paragraph_format.space_after = Pt(28)

    def format_abstract_in_english(doc, text_to_format):
        for paragraph in doc.paragraphs:
            if text_to_format in paragraph.text:
                for run in paragraph.runs:
                    # 设置字体为Times New Roman
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    # 设置字号为三号字（16磅）
                    run.font.size = Pt(16)
                    # 设置字体加粗
                    run.font.bold = True
                    # 设置居中对齐
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 设置单倍行距
                    paragraph.paragraph_format.line_spacing = 1
                    # 设置段前2行（约28磅）
                    paragraph.paragraph_format.space_before = Pt(28)
                    # 设置段后2行（约28磅）
                    paragraph.paragraph_format.space_after = Pt(28)

    def set_single_line_spacing_for_images(doc):
        """遍历文档中的所有段落，如果包含图片，则设置单倍行距。"""
        for paragraph in doc.paragraphs:
            if has_graphics(paragraph):
                # 设置段落的行距为单倍行距
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def format_normal_text_in_document(doc, start_paragraph_index):
        skip = False  # 跳过标志初始化为False
        for paragraph in doc.paragraphs[start_paragraph_index:]:
            if paragraph.style.name == 'Heading 1':
                if "参考文献" in paragraph.text:
                    skip = True  # 开始跳过
                elif "致谢" in paragraph.text:
                    skip = False  # 停止跳过
                    continue  # 确保“致谢”之后的段落不被跳过

            if skip:
                continue  # 如果处于跳过状态，忽略当前段落的处理

            if paragraph.style.name == 'Normal':
                if "摘    要" not in paragraph.text and "ABSTRACT" not in paragraph.text:
                    if "目    录" not in paragraph.text:
                        if not has_graphics(paragraph):
                            # print("Processing paragraph:", paragraph.text)
                            for run in paragraph.runs:
                                if not run.font.superscript:
                                    # 设置中文字符为宋体小四号字
                                    run.font.name = '宋体'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                    # 设置英文字符为Times New Roman小四号字
                                    run.font.name = 'Times New Roman'
                                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                                    # 设置字号为小四号字（12磅）
                                    run.font.size = Pt(12)
                                    # 设置行距为固定值20磅
                                    paragraph.paragraph_format.line_spacing = Pt(20)
                                    if paragraph.paragraph_format.first_line_indent is None:
                                        paragraph.paragraph_format.first_line_indent = Cm(0.85)
                                        paragraph.paragraph_format.left_indent = Cm(0)
                                    else:
                                        if paragraph.paragraph_format.first_line_indent > Cm(0):
                                            paragraph.paragraph_format.first_line_indent = Cm(0.85)
                                            paragraph.paragraph_format.left_indent = Cm(0)

    # def set_font_according_to_language(run):
    #     # 根据运行的文本设置字体，中文使用宋体，英文使用Times New Roman
    #     if re.search('[\u4e00-\u9fff]', run.text):
    #         run.font.name = '宋体'
    #         run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    #     else:
    #         run.font.name = 'Times New Roman'
    #         run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    #         run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    #
    # def format_normal_text_in_document(doc, start_paragraph_index):
    #     skip = False  # 跳过标志初始化为False
    #     for paragraph in doc.paragraphs[start_paragraph_index:]:
    #         if paragraph.style.name == 'Heading 1':
    #             if "参考文献" in paragraph.text:
    #                 skip = True  # 开始跳过
    #             elif "致谢" in paragraph.text:
    #                 skip = False  # 停止跳过
    #                 continue  # 确保“致谢”之后的段落不被跳过
    #
    #         if skip:
    #             continue  # 如果处于跳过状态，忽略当前段落的处理
    #
    #         if paragraph.style.name == 'Normal':
    #             if paragraph.style.name == 'Normal' and "摘    要" not in paragraph.text and "ABSTRACT" not in paragraph.text:
    #                 if paragraph.style.name == 'Normal' and "目    录" not in paragraph.text:
    #                     if not has_graphics(paragraph):
    #                         # print("Processing paragraph:", paragraph.text)
    #                         for run in paragraph.runs:
    #                             set_font_according_to_language(run)
    #                             run.font.size = Pt(12)  # 设置字号为小四号字（12磅）
    #                         paragraph.paragraph_format.line_spacing = Pt(20)  # 设置行距为固定值20磅

    def add_page_break(doc, text_to_format):
        for i, paragraph in enumerate(doc.paragraphs):
            if text_to_format in paragraph.text:
                # 在包含“ABSTRACT”的段落前添加分页符
                if i > 0:  # 确保不是文档的第一个段落
                    doc.paragraphs[i - 1].add_run().add_break()
                break  # 找到“ABSTRACT”后退出循环

    def modify_and_format_text_on_first_page(doc, modifications, formats):
        first_page_paragraphs = doc.paragraphs[:7]  # 假设第一页的内容在前10个段落中
        for paragraph in first_page_paragraphs:
            # print("Processing paragraph:", paragraph.text)
            for old_text, new_text in modifications.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
            for text, format_settings in formats.items():
                font, size, bold = format_settings  # 解包格式设置
                if text in paragraph.text:
                    for run in paragraph.runs:
                        if text in run.text:
                            # 设置字体
                            run.font.name = font
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
                            # 设置字号
                            run.font.size = Pt(size)
                            # 设置加粗
                            run.font.bold = bold

    def set_page_break_before(paragraph):
        """为指定段落设置段前分页"""
        p = paragraph._p  # 访问底层的xml元素
        pPr = p.get_or_add_pPr()  # 获取或添加段落属性
        pageBreakBefore = OxmlElement('w:pageBreakBefore')  # 创建段前分页元素
        pPr.append(pageBreakBefore)  # 将段前分页添加到段落属性中

    def has_page_break_before(paragraph):
        """检查段落前是否有分页符。"""
        # 检查该段落之前的所有run，看看是否包含分页符
        prev = paragraph._element.getprevious()
        while prev is not None:
            if prev.tag.endswith('p'):
                for r in prev.findall('.//w:r', namespaces=nsmap):
                    for br in r.findall('.//w:br', namespaces=nsmap):
                        if br.get(qn('w:type')) == 'page':
                            return True
            prev = prev.getprevious()
        return False

    def apply_style_settings(paragraph, font_name, font_size, alignment, line_spacing, space_before, space_after):
        # 重新应用样式以重置所有属性
        paragraph.style = paragraph.style
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
        paragraph.alignment = alignment
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)

    def format_headings_in_document(doc):
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                # print("Processing paragraph:", paragraph.text)
                paragraph.style = doc.styles['Heading 1']
                apply_style_settings(paragraph, '黑体', 16, WD_ALIGN_PARAGRAPH.CENTER, 1, 28, 28)
            elif paragraph.style.name == 'Heading 2':
                # print("Processing paragraph:", paragraph.text)
                paragraph.style = doc.styles['Heading 2']
                apply_style_settings(paragraph, '黑体', 14, WD_ALIGN_PARAGRAPH.LEFT, Pt(20), 0, 0)
            elif paragraph.style.name == 'Heading 3':
                # print("Processing paragraph:", paragraph.text)
                paragraph.style = doc.styles['Heading 3']
                apply_style_settings(paragraph, '黑体', 12, WD_ALIGN_PARAGRAPH.LEFT, Pt(20), 0, 0)

    def format_keywords(doc):
        toc_found = False  # 用于标记是否找到“目    录”
        for paragraph in doc.paragraphs:
            # print("Processing paragraph:", paragraph.text)
            if "目    录" in paragraph.text:
                toc_found = True
                break
            if not has_graphics(paragraph):
                # 处理中文关键词
                if "关键词：" in paragraph.text:
                    process_keywords(paragraph, "关键词：", '黑体', '宋体', False)
                    paragraph.paragraph_format.space_after = Pt(0)
                elif "Key words：" in paragraph.text:
                    process_keywords(paragraph, "Key words：", 'Times New Roman', 'Times New Roman', True)
                    paragraph.paragraph_format.space_after = Pt(0)
            if toc_found:
                break  # 如果已经处理到“目 录”，则停止进一步处理

    def process_keywords(paragraph, keyword_text, keyword_font, text_font, bold):
        # 拆分原始文本以提取和格式化关键词部分
        parts = paragraph.text.split(keyword_text)
        paragraph.clear()

        # 添加关键词前的文本，假设前文使用宋体
        if parts[0].strip():
            add_text_run(paragraph, parts[0], text_font, False)

        # 添加格式化的关键词，使用黑体或Times New Roman
        add_text_run(paragraph, keyword_text, keyword_font, bold)

        # 添加关键词后的文本，使用宋体或Times New Roman
        if len(parts) > 1:
            add_text_run(paragraph, parts[1], text_font, False)

    def add_text_run(paragraph, text, font_name, bold):
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia' if font_name == '黑体' else 'w:ascii'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi' if font_name == 'Times New Roman' else 'w:eastAsia'), font_name)
        run.font.size = Pt(12)  # 设置字号为12磅

    def determine_font(text):
        """根据文本内容决定使用的字体。"""
        if any('\u4e00' <= char <= '\u9fff' for char in text):
            return '宋体'
        else:
            return 'Times New Roman'

    def should_add_page_break(paragraphs, index):
        # 检查前1个段落是否有文本，返回True如果有任何一个包含文本
        start_index = max(index - 3, 0)
        for i in range(start_index, index):
            if paragraphs[i].text.strip():
                return True
        return False

    def add_page_break_before_headings(doc):
        paragraphs = list(doc.paragraphs)
        first_heading_found = False
        for i, paragraph in enumerate(paragraphs):
            if paragraph.style.name == 'Heading 1':
                # print("Processing paragraph:", paragraph.text)
                if not first_heading_found:
                    first_heading_found = True
                    continue  # 跳过文档中的第一个一级标题
                if should_add_page_break(paragraphs, i):
                    set_page_break_before(paragraph)

    def remove_blank_heading_ones(doc):
        # 准备一个列表来收集需要删除的段落的引用
        paragraphs_to_remove = []

        # 遍历所有段落，找到空白的一级标题
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1' and not paragraph.text.strip():
                # 添加到待删除列表
                paragraphs_to_remove.append(paragraph)

        # 从文档中删除收集的空白段落
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)

    def modify_figure_paragraphs(doc):
        # 正则表达式匹配以“图”开头，后面跟数字的段落
        figure_pattern = r'^(表|图)\d+.*$'

        for paragraph in doc.paragraphs:
            if re.match(figure_pattern, paragraph.text):
                # 设置段落居中对齐
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 对中文和英文字符设置不同字体和字号
                for run in paragraph.runs:
                    # 取消加粗
                    run.font.bold = False
                    if any(char in run.text for char in
                           'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890-'):
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10.5)  # 五号字大约等于10.5磅
                        # 设置固定行距为20磅
                        paragraph.paragraph_format.line_spacing = Pt(20)
                    else:
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)  # 五号字大约等于10.5磅
                        # 设置固定行距为20磅
                        paragraph.paragraph_format.line_spacing = Pt(20)

    def operate_cited(doc):
        start_index, end_index = None, None

        # 找到“参考文献”和“致谢”的索引
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text == "参考文献" and paragraph.style.name == 'Heading 1':
                start_index = i + 1  # 从"参考文献"之后的段落开始处理
            elif text == "致谢" and paragraph.style.name == 'Heading 1':
                end_index = i  # 处理直到"致谢"之前的段落
        if start_index is not None and end_index is not None and start_index < end_index:
            for paragraph in doc.paragraphs[start_index:end_index]:
                match = re.search(r'(\D+)(\d{1,2})(\D+)', paragraph.text)
                if match:
                    before, number, after = match.groups()
                    new_text = before.replace(before[-1], '[') + number + after.replace(after[0], ']')
                    paragraph.text = re.sub(r'(\D+)(\d{1,2})(\D+)', new_text, paragraph.text, count=1)
                for run in paragraph.runs:
                    if run.text.strip() != '':
                        if any(ord(c) > 128 for c in run.text):  # 判断是否包含中文字符
                            run.font.name = '宋体'
                        else:
                            run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph.paragraph_format.line_spacing = Pt(20)

                indent_xml = """
                           <w:ind xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:left="0" w:hanging="280"/>
                       """
                indent_element = parse_xml(indent_xml)
                paragraph._p.insert(2, indent_element)
                modified_text = paragraph.text.replace('，', ',')
                modified_text = modified_text.replace('。', '.')
                modified_text = modified_text.replace('；', ';')
                modified_text = modified_text.replace('：', ':')
                modified_text = modified_text.replace('【', '[')
                modified_text = modified_text.replace('】', ']')
                modified_text = modified_text.replace('（', '(')
                modified_text = modified_text.replace('）', ')')
                modified_text = modified_text.replace('. ', '.')
                paragraph.text = modified_text

    def set_normal_style_between_sections(doc):
        in_abstract_to_toc = False  # 控制是否在“摘要”到“目录”之间
        in_main_content = False  # 控制是否在主要内容部分（第一个一级标题之后）
        toc_found = False  # 标记是否找到了“目录”

        for paragraph in doc.paragraphs:
            # 标记“摘要”开始
            if "摘    要" in paragraph.text:
                in_abstract_to_toc = True

            # 标记“目录”找到，并结束“摘要”到“目录”的部分
            if "目    录" in paragraph.text:
                in_abstract_to_toc = False
                toc_found = True

            # 如果找到“目录”且当前段落是一级标题，则开始主要内容部分
            if toc_found and paragraph.style.name == 'Heading 1':
                in_main_content = True
                toc_found = False  # 重置，以防之后有其他内容需要处理

            # 如果在“摘要”到“目录”之间，或者在主要内容部分
            if in_abstract_to_toc or (in_main_content and not paragraph.style.name.startswith('Heading')):
                paragraph.style = doc.styles['Normal']

                # for paragraph in doc.paragraphs:
                #     if paragraph.style.name == '正文':
                #         # 获取段落属性
                #         pPr = paragraph._p.get_or_add_pPr()
                #         # 设置段前距
                #         spacing_before = OxmlElement('w:spacing')
                #         spacing_before.set(qn('w:before'), "0")
                #         # 设置段后距
                #         spacing_after = OxmlElement('w:spacing')
                #         spacing_after.set(qn('w:after'), "0")
                #         # 将设置添加到段落属性中
                #         pPr.append(spacing_before)
                #         pPr.append(spacing_after)

    # def format_specific_keywords_and_text_after_them(doc):
    #     keywords = ["题    目", "学院名称", "专    业", "班    级", "学    号", "学生姓名", "指导教师", "完成日期"]
    #     abstract_found = False  # 标记是否找到了“摘要”
    #
    #     for paragraph in doc.paragraphs:
    #         # 如果找到“摘要”，停止格式化
    #         if "摘    要" in paragraph.text:
    #             abstract_found = True
    #             break
    #
    #         original_text = paragraph.text
    #         for keyword in keywords:
    #             if keyword in original_text:
    #                 # 找到关键字，分割文本
    #                 parts = original_text.split(keyword, 1)
    #                 if len(parts) > 1:
    #                     # 清除原段落内容
    #                     paragraph.clear()
    #                     # 重写并格式化关键字部分（宋体小三号，加粗）
    #                     run = paragraph.add_run(parts[0] + keyword)
    #                     run.font.name = '宋体'
    #                     run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    #                     run.font.size = Pt(16)  # 设置字号为小三号（约16磅）
    #                     run.font.bold = True
    #
    #                     # 重写并格式化关键字之后的部分（楷体小三号）
    #                     run = paragraph.add_run(parts[1])
    #                     run.font.name = 'KaiTi'  # 设置字体为楷体
    #                     run._element.rPr.rFonts.set(qn('w:eastAsia'), 'KaiTi')
    #                     run.font.size = Pt(15)  # 设置字号为小三号（约15磅）
    #                     run.font.bold = False

    # def operate_normal_symbol(doc):
    #     half_to_full_map = {
    #         ',': '，',
    #         '?': '？',
    #         '!': '！',
    #         ':': '：',
    #         ';': '；',
    #         '(': '（',
    #         ')': '）',
    #         '{': '｛',
    #         '}': '｝'
    #     }
    #
    #     start_processing = False  # 使用一个标志来控制处理的开始
    #     for paragraph in doc.paragraphs:
    #         if paragraph.style.name == 'Heading 1' and '1  引言' in paragraph.text:
    #             start_processing = True  # 开始处理文档
    #         if not paragraph.style.name == 'Heading':
    #             if start_processing and paragraph.style.name == 'Normal':
    #                 if not has_graphics(paragraph):
    #                     for run in paragraph.runs:
    #                         new_text = ''
    #                         for char in run.text:
    #                             if char in half_to_full_map:
    #                                 new_text += half_to_full_map[char]
    #                             else:
    #                                 new_text += char
    #                         run.text = new_text

    def operate_normal_symbol(doc):
        half_to_full_map = {
            ',': '，',
            '?': '？',
            '!': '！',
            ':': '：',
            ';': '；',
            '(': '（',
            ')': '）',
            '{': '｛',
            '}': '｝'
        }

        punctuation_marks = set("。？！，、；：「」『』（）《》【】〈〉“”‘’…—")
        brackets = {'（', '）', '(', ')'}  # 定义括号集合，可根据需要添加全角和半角的其它括号类型

        start_processing = False  # 使用一个标志来控制处理的开始

        pattern = r'.*引言$'

        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                if re.fullmatch(pattern, paragraph.text.strip()):
                    start_processing = True  # 开始处理文档

            if start_processing and paragraph.style.name == 'Normal':
                if not has_graphics(paragraph):  # 确保不处理包含图形的段落
                    for run in paragraph.runs:
                        # 转换半角到全角符号
                        new_text = ''.join([half_to_full_map.get(char, char) for char in run.text])
                        run.text = new_text

                    # 删除连续的标点符号中的第一个
                    original_text = paragraph.text
                    new_text = []
                    chars = list(original_text)  # 转换为字符列表以处理连续标点
                    i = 0
                    # while i < len(chars):
                    #     if i + 1 < len(chars) and chars[i] in punctuation_marks and chars[i + 1] in punctuation_marks:
                    #         i += 1  # 跳过第一个标点，保留第二个
                    #     new_text.append(chars[i])
                    #     i += 1

                    while i < len(chars) - 1:  # 减1避免索引越界
                        # 检查当前和下一个字符是否都在标点集合中，且不含括号
                        if chars[i] in punctuation_marks and chars[i + 1] in punctuation_marks and \
                                (chars[i] not in brackets and chars[i + 1] not in brackets):
                            # 删除当前的重复标点符号，保留下一个（通过不增加i来间接实现）
                            del chars[i]
                        else:
                            i += 1  # 移动到下一个字符

                    paragraph.text = ''.join(chars)  # 更新段落文本

    def set_continuous_heading_numbers(doc_path):
        heading_levels = {}  # 用于存储各级标题的当前计数
        unnumbered_headings = ["参考文献", "致谢"]  # 不进行编号的特定标题列表

        for paragraph in doc.paragraphs:
            # 检查段落是否为标题
            if paragraph.style.name.startswith('Heading'):
                # 跳过不编号的特定标题
                if any(unnumbered_heading in paragraph.text for unnumbered_heading in unnumbered_headings):
                    continue

                level = int(paragraph.style.name.split(' ')[1])  # 获取标题级别
                # 初始化或更新标题级别计数
                if level in heading_levels:
                    # 重置下一级标题的计数
                    if level == 1:
                        heading_levels[2] = 0  # 重置二级标题计数
                        heading_levels[3] = 0  # 重置三级标题计数
                    elif level == 2:
                        heading_levels[3] = 0  # 重置三级标题计数
                    heading_levels[level] += 1
                else:
                    # 如果是新的标题级别，重置该级别和所有更低级别的计数
                    for l in range(level, 10):  # 假设标题不会超过9级
                        heading_levels[l] = 0
                    heading_levels[level] = 1

                # 生成新的标题编号
                number = '.'.join(
                    str(heading_levels[l]) for l in sorted(heading_levels) if heading_levels[l] > 0 and l <= level)
                # 更新段落文本以包含新的标题编号
                if paragraph.text.startswith(number):
                    # 如果已有正确的编号，跳过
                    continue
                else:
                    # 移除原有编号
                    text_without_number = paragraph.text.split(' ', 1)[-1] if ' ' in paragraph.text else paragraph.text
                    paragraph.text = f"{number} {text_without_number}"

    def align_paragraphs_left(doc_path):
        keywords = ['关键词', 'Key words']
        paragraphs_to_check = list(doc.paragraphs)  # 创建段落列表的副本

        for i in range(1, len(paragraphs_to_check)):  # 从1开始以确保i-1有效
            paragraph = paragraphs_to_check[i]
            if any(keyword in paragraph.text for keyword in keywords):
                previous_paragraph = paragraphs_to_check[i - 1]
                if previous_paragraph and previous_paragraph.text.strip():
                    # 在关键词段落前插入一个新的空白段落
                    new_paragraph = paragraph.insert_paragraph_before()  # 正确插入段落
                    new_paragraph.text = ''  # 确保新段落为空
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = None

    def split_keywords(doc_path):

        keywords_chinese = '关键词'
        keywords_english = 'Key words'

        for paragraph in doc.paragraphs:
            # 检查是否包含中文关键词标识或英文关键词标识
            if keywords_chinese in paragraph.text or keywords_english in paragraph.text:
                for run in paragraph.runs:
                    if '、' in run.text:
                        run.text = run.text.replace('、', '；')
                    if ',' in run.text:
                        run.text = run.text.replace(',', '；')
                    if ':' in run.text:
                        run.text = run.text.replace(':', '：')
                    if ';' in run.text:
                        run.text = run.text.replace(';', '；')

    def set_heading_one_for_specific_paragraphs(doc_path):
        target_texts = ["致谢", "参考文献"]  # 指定需要修改为一级标题的段落文本

        for paragraph in doc.paragraphs:
            if paragraph.text.strip() in target_texts:  # 检查段落文本是否精确匹配列表中的任何一个元素
                # print("Processing paragraph:", paragraph.text)
                paragraph.style = doc.styles['Heading 1']  # 应用一级标题样式
                # 寻找并删除编号相关的XML节点
                pPr = paragraph._element.pPr
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)

    # def remove_numbering_from_specific_headings(doc_path):
    #     doc = Document(doc_path)
    #     target_titles = ["参考文献", "致谢"]
    #
    #     for paragraph in doc.paragraphs:
    #         if paragraph.text.strip() in target_titles:
    #             # 寻找并删除编号相关的XML节点
    #             pPr = paragraph._element.pPr
    #             numPr = pPr.find(qn('w:numPr'))
    #             if numPr is not None:
    #                 pPr.remove(numPr)

    # print("Processing paragraph:", paragraph.text)

    def modify_table(doc_path):
        table = doc.tables[0]

        # 遍历表格中的行
        for row in table.rows:
            # 获取第二列的单元格
            cell = row.cells[1]
            # 遍历单元格中的所有段落，设置字体和大小
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
                for run in paragraph.runs:
                    run.font.name = '楷体'  # 设置字体为楷体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'KaiTi')
                    run.font.size = Pt(15)  # 设置字体大小为小三号大小

    def remove_numbering(paragraph):
        """Remove numbering from a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        numPr = pPr.numPr
        if numPr:
            pPr.remove(numPr)

    def add_numbering_between_headings(doc):
        start_numbering = False
        counter = 1  # 编号开始于1
        heading_style = 'Heading 1'  # 一级标题的样式名
        reference_heading_found = False
        pattern = re.compile(r'^\[\d+\]\s*|\d+\.\s*|\(\d+\)\s*')  # 正则表达式识别编号

        for paragraph in doc.paragraphs:
            if '参考文献' in paragraph.text and paragraph.style.name == heading_style:
                start_numbering = True
                reference_heading_found = True
                continue

            if '致谢' in paragraph.text and paragraph.style.name == heading_style and reference_heading_found:
                break

            if start_numbering and reference_heading_found:
                if paragraph.style.name != heading_style and paragraph.text.strip():
                    original_text = pattern.sub('', paragraph.text.strip())
                    remove_numbering(paragraph)  # 移除自动编号
                    paragraph.clear()
                    paragraph.add_run(f"[{counter}] {original_text}")
                    counter += 1
                # print(f"编号 {counter - 1}: {paragraph.text[:30]}")  # 显示编号后的段落前30个字符

                # 设置行间距
                paragraph.paragraph_format.line_spacing = Pt(20)
                paragraph.paragraph_format.first_line_indent = -Cm(0.74)  # 每个字符大约0.1英寸
                paragraph.paragraph_format.left_indent = Cm(0.74)

    def add_footer_with_auto_numbering(doc_path):

        word = win32.gencache.EnsureDispatch('Word.Application')
        doc = word.Documents.Open(doc_path)
        word.Visible = False  # 设置为 True 可以看到操作过程，便于调试
        # 查找第一个“标题 1”并插入分节符
        found = False
        for paragraph in doc.Paragraphs:
            if paragraph.Style.NameLocal == '标题 1':  # 根据你的Word版本调整样式名称
                found = True
                # 在“标题 1”起始处插入分节符
                range = paragraph.Range
                range.Collapse(Direction=1)  # Collapse the range to its start
                range.InsertBreak(Type=win32.constants.wdSectionBreakNextPage)
                break
        if found:
            # 获取最新创建的分节
            section = doc.Sections.Last
            # 设置新分节的页脚
            footer = section.Footers(win32.constants.wdHeaderFooterPrimary)
            footer.LinkToPrevious = False  # 不与前一节的页脚链接

            # 重置页码开始为 1
            footer.PageNumbers.RestartNumberingAtSection = True
            footer.PageNumbers.StartingNumber = 1  # 正确设置页码从 1 开始

            # 插入格式化的页码字段，并确保破折号正确插入
            footer_text = footer.Range
            footer_text.Text = " - "  # 首先插入前破折号
            footer_text.Collapse(Direction=1)  # Collapse range to the end
            footer_text.Fields.Add(footer_text, win32.constants.wdFieldEmpty, r'PAGE \* Arabic \* MERGEFORMAT', True)
            footer_text.InsertAfter(" - ")  # 在页码后添加第二个破折号
            footer.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
            # 保存更改
            doc.Save()
        # 关闭文档和 Word 应用
        doc.Close()
        word.Quit()

    def set_font_size(paragraph, size):
        for run in paragraph.runs:
            run.font.size = Pt(size)

    def renum(doc_path):
        def is_numbered_paragraph(paragraph):
            # 匹配例如 "（1） ", "（2） ", "（10） " 等格式
            return re.match(r'^（\d+）', paragraph.text)

        # 初始化编号
        current_number = 1

        # 重新编号带有全角括号编号的段落
        for para in doc.paragraphs:
            if is_numbered_paragraph(para):
                # 提取当前编号
                match = re.match(r'^（(\d+)）', para.text)
                if match:
                    original_number = int(match.group(1))
                    if original_number == 1:
                        current_number = 1  # 遇到 "（1）" 则重新开始编号
                    else:
                        # 只对非 "（1）" 的段落重新编号
                        if current_number == 1:
                            current_number += 1
                        para.text = re.sub(r'^（\d+）', f'（{current_number}）', para.text)
                        current_number += 1
                    # 设置段落字体为小四号（12磅）
                    set_font_size(para, 12)

    def format_normal_text_in_abstract(doc, start_paragraph_index):
        skip = False
        for paragraph in doc.paragraphs[start_paragraph_index:]:
            if "摘    要" in paragraph.text:
                skip = False
            elif "目    录" in paragraph.text:
                skip = True

            if skip:
                break

            if paragraph.style.name == 'Normal':
                if paragraph.style.name == 'Normal' and "摘    要" not in paragraph.text and "ABSTRACT" not in paragraph.text:
                    if paragraph.style.name == 'Normal' and "目    录" not in paragraph.text:
                        # print("Processing paragraph:", paragraph.text)
                        for run in paragraph.runs:
                            if not run.font.superscript:
                                # 设置中文字符为宋体小四号字
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                # 设置英文字符为Times New Roman小四号字
                                run.font.name = 'Times New Roman'
                                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                                # 设置字号为小四号字（12磅）
                                run.font.size = Pt(12)
                                # 设置行距为固定值20磅
                                paragraph.paragraph_format.line_spacing = Pt(20)
                                if paragraph.paragraph_format.first_line_indent is None:
                                    paragraph.paragraph_format.first_line_indent = Cm(0.85)
                                    paragraph.paragraph_format.left_indent = Cm(0)
                                else:
                                    if paragraph.paragraph_format.first_line_indent > Cm(0):
                                        paragraph.paragraph_format.first_line_indent = Cm(0.85)
                                        paragraph.paragraph_format.left_indent = Cm(0)

    def update_paragraphs(doc):
        # 定义用于匹配 "图X-X *" 的正则表达式
        pattern = re.compile(r'^图(\d+)-(\d+)(.*)')

        # 用于记录每个第一个X的最新第二个X值
        figure_count = {}

        # 遍历所有段落
        for para in doc.paragraphs:
            new_text = para.text
            # 查找所有匹配的文本
            matches = pattern.findall(new_text)

            if matches:
                # 初始化临时文本
                temp_text = new_text
                for match in matches:
                    first_x_value = int(match[0])
                    second_x_value = int(match[1])
                    rest_of_text = match[2]

                    # 如果是新的第一个X，初始化计数为1
                    if first_x_value not in figure_count:
                        figure_count[first_x_value] = 1

                    # 获取应当替换的第二个X值
                    expected_x_value = figure_count[first_x_value]

                    # 构建旧的和新的字符串
                    old_str = f'图{first_x_value}-{second_x_value}{rest_of_text}'
                    new_str = f'图{first_x_value}-{expected_x_value}{rest_of_text}'

                    # 只替换第一个匹配的字符串
                    temp_text = temp_text.replace(old_str, new_str, 1)

                    # 更新计数
                    figure_count[first_x_value] += 1

                # 更新段落文本
                para.text = temp_text

    selected_doc_path = select_document()
    if selected_doc_path:
        doc = Document(selected_doc_path)

        messagebox.showinfo("Info", "处理文档需要一些时间，请耐心等待")

        # 格式化表格
        modify_table(doc)

        # 删除所有空白的一级标题
        remove_blank_heading_ones(doc)
        # remove_blank_paragraphs(doc)

        # 格式化文献格式
        operate_cited(doc)

        # 调整标题顺序
        set_heading_one_for_specific_paragraphs(doc)
        set_continuous_heading_numbers(doc)

        add_numbering_between_headings(doc)

        update_paragraphs(doc)

        # 设置文档段落为Normal样式
        set_normal_style_between_sections(doc)

        # 设置布局
        set_page_layout(doc)

        header_text_one = "杭州电子科技大学信息工程学院本科毕业设计"
        # 格式化页眉
        header_text = header_text_one
        update_headers_if_text_exists(doc, header_text)

        format_normal_text_in_abstract(doc, 5)

        # 需要修改的文本，字体和字号
        modifications = {
            "杭州电子科技大学信息工程学院": "杭州电子科技大学信息工程学院",
            "本科毕业设计": "本科毕业设计",
            "（2024届）": "（2024届）"
        }
        # 要设置的文本，字体和字号
        formats = {
            "杭州电子科技大学信息工程学院": ("宋体", 28, True),
            "本科毕业设计": ("宋体", 28, True),
            "（2024届）": ("宋体", 22, True)  # 二号字号一般为22磅
        }

        # 修改并格式化第一页中的特定文本
        modify_and_format_text_on_first_page(doc, modifications, formats)

        # 摘要修改
        text_to_format1 = "摘    要"
        format_abstract(doc, text_to_format1)
        text_to_format2 = "ABSTRACT"
        format_abstract_in_english(doc, text_to_format2)
        text_to_format3 = "目    录"
        format_abstract(doc, text_to_format3)

        # 操作标点
        operate_normal_symbol(doc)

        # 格式化正文内容
        format_normal_text_in_document(doc, 6)

        # 更新标题样式
        format_headings_in_document(doc)

        # 修改图注
        modify_figure_paragraphs(doc)

        # 图片设置为单倍行距
        set_single_line_spacing_for_images(doc)

        # 格式化关键词
        format_keywords(doc)
        align_paragraphs_left(doc)
        split_keywords(doc)

        # 为符合条件的一级标题添加段前分页
        add_page_break_before_headings(doc)

        renum(doc)

        # 另存为文件
        new_doc_path = select_save_as()
        if new_doc_path:
            doc.save(new_doc_path)
            add_footer_with_auto_numbering(new_doc_path)
            messagebox.showinfo("Info", "文件已处理完毕")
        else:
            messagebox.showinfo("Info", "取消保存文件")
    else:
        messagebox.showinfo("Info", "未选择文件或者取消")


def apply_custom_template_1():
    root = tk.Tk()
    root.withdraw()  # 隐藏主窗口

    def set_normal_style_between_sections(doc):
        in_abstract_to_toc = False  # 控制是否在“摘要”到“目录”之间
        in_main_content = False  # 控制是否在主要内容部分（第一个一级标题之后）
        toc_found = False  # 标记是否找到了“目录”

        for paragraph in doc.paragraphs:
            # 标记“摘要”开始
            if "摘    要" in paragraph.text:
                in_abstract_to_toc = True

            # 标记“目录”找到，并结束“摘要”到“目录”的部分
            if "目    录" in paragraph.text:
                in_abstract_to_toc = False
                toc_found = True

            # 如果找到“目录”且当前段落是一级标题，则开始主要内容部分
            if toc_found and paragraph.style.name == 'Heading 1':
                in_main_content = True
                toc_found = False  # 重置，以防之后有其他内容需要处理

            # 如果在“摘要”到“目录”之间，或者在主要内容部分
            if in_abstract_to_toc or (in_main_content and not paragraph.style.name.startswith('Heading')):
                paragraph.style = doc.styles['Normal']

    def get_user_input():
        root = tk.Tk()
        root.withdraw()  # 隐藏主窗口

        # 获取系统字体列表
        system_fonts = font.families(root)

        # 创建一个新窗口来选择字体
        font_selector = tk.Toplevel(root)
        font_selector.title("自定义标题样式")
        font_selector.geometry("300x250")

        def create_font_widgets(level, font_list):
            tk.Label(font_selector, text=f"请输入{level}标题的字体名称和字号:").pack()
            font_var = tk.StringVar(font_selector)
            font_var.set(font_list[0])  # 默认设置为列表中的第一个字体
            font_menu = tk.OptionMenu(font_selector, font_var, *font_list)
            font_menu.pack()

            font_size = simpledialog.askinteger("输入字号", f"请输入{level}标题的字号 (pt):", parent=font_selector)

            return font_var, font_size

        font_name_h1, font_size_h1 = create_font_widgets("一级", system_fonts)
        font_name_h2, font_size_h2 = create_font_widgets("二级", system_fonts)
        font_name_h3, font_size_h3 = create_font_widgets("三级", system_fonts)

        def on_submit():
            styles = {
                'h1': {'font_name': font_name_h1.get(), 'font_size': font_size_h1},
                'h2': {'font_name': font_name_h2.get(), 'font_size': font_size_h2},
                'h3': {'font_name': font_name_h3.get(), 'font_size': font_size_h3}
            }
            font_selector.destroy()  # 销毁字体选择器窗口
            root.quit()  # 退出主事件循环
            return styles

        submit_button = tk.Button(font_selector, text="确认", command=on_submit)
        submit_button.pack()

        root.mainloop()
        return on_submit()

    def apply_style_settings_cus(paragraph, font_name, font_size, alignment, line_spacing, space_before, space_after):
        for run in paragraph.runs:
            # 设置正文和复杂脚本字体
            run.font.name = font_name
            run.font.cs = font_name  # 设置复杂脚本字体，如果需要
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 设置东亚字体

            run.font.size = Pt(font_size)

        paragraph.alignment = alignment
        paragraph.paragraph_format.line_spacing = Pt(line_spacing)
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)

        # 确保至少有一个运行使用这个样式
        if not paragraph.runs:
            run = paragraph.add_run("")
            run.font.name = font_name
            run.font.cs = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)

    def format_document(doc_path, styles):
        doc = Document(doc_path)
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                apply_style_settings_cus(paragraph, styles['h1']['font_name'], styles['h1']['font_size'],
                                         WD_ALIGN_PARAGRAPH.CENTER, 1, 28, 28)
            elif paragraph.style.name == 'Heading 2':
                apply_style_settings_cus(paragraph, styles['h2']['font_name'], styles['h2']['font_size'],
                                         WD_ALIGN_PARAGRAPH.LEFT, 20, 0, 0)
            elif paragraph.style.name == 'Heading 3':
                apply_style_settings_cus(paragraph, styles['h3']['font_name'], styles['h3']['font_size'],
                                         WD_ALIGN_PARAGRAPH.LEFT, 20, 0, 0)
        return doc

    def ttd():
        doc_path = select_document()
        if not doc_path:
            messagebox.showinfo("Info", "未选择文档，操作取消。")
            return

        # 获取输入
        styles = get_user_input()
        # 应用修改
        doc = format_document(doc_path, styles)
        set_normal_style_between_sections(doc)

        save_path = select_save_as()
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("Info", "文件已处理完毕")
        else:
            messagebox.showinfo("Info", "未选择保存位置，操作取消。")

    ttd()
    root.mainloop()


def apply_custom_template_2():
    def set_page_layout_cus(doc, layout):
        for section in doc.sections:
            section.page_width = Mm(layout['页面宽度 (mm)'])
            section.page_height = Mm(layout['页面高度 (mm)'])
            section.top_margin = Mm(layout['上边距 (mm)'])
            section.bottom_margin = Mm(layout['下边距 (mm)'])
            section.left_margin = Mm(layout['左边距 (mm)'])
            section.right_margin = Mm(layout['右边距 (mm)'])
            section.gutter = Mm(layout['装订线 (mm)'])
            section.footer_distance = Cm(layout['页脚距离 (cm)'])
            section.header_distance = Cm(layout['页眉距离 (cm)'])

    def get_layout_input(doc):
        root = tk.Tk()
        root.title("设置页面布局")

        entries = {}
        # 这里定义的标签应与下面创建字典时使用的键匹配
        labels = ['页面宽度 (mm)', '页面高度 (mm)', '上边距 (mm)', '下边距 (mm)',
                  '左边距 (mm)', '右边距 (mm)', '装订线 (mm)', '页脚距离 (cm)', '页眉距离 (cm)']
        defaults = [210, 297, 30, 20, 30, 20, 10, 1, 2]

        for label, default in zip(labels, defaults):
            frame = tk.Frame(root)
            frame.pack(fill='x', padx=5, pady=5)
            tk.Label(frame, text=label).pack(side='left')
            entry_var = StringVar(value=str(default))
            entry = tk.Entry(frame, textvariable=entry_var)
            entry.pack(side='left')
            entries[label] = entry

        def submit2():
            # 确保这里的字典键与 set_page_layout_cus 函数中的参数名一致
            layout = {label: float(entries[label].get()) for label in labels}
            set_page_layout_cus(doc, layout)
            save_path = select_save_as()
            doc.save(save_path)
            messagebox.showinfo("Info", "文件已处理完毕")
            root.destroy()

        tk.Button(root, text="确认", command=submit2).pack(pady=10)
        root.mainloop()

    doc_path = select_document()
    if doc_path:
        doc = Document(doc_path)
        get_layout_input(doc)
    else:
        messagebox.showinfo("Info", "未选择文档，操作取消。")


def apply_custom_template_3():
    def get_font_settings(root, callback):
        # 获取系统中的字体列表
        system_fonts = families(root)

        # 设置选择字体的对话框
        font_choice = tk.Toplevel(root)
        font_choice.title("设置字体和字号")
        font_choice.geometry("250x180")

        # 中文字体选择
        chinese_font_var = StringVar(font_choice)
        chinese_font_var.set(system_fonts[0])  # 默认设置为第一个字体
        Label(font_choice, text="选择中文字体:").pack()
        OptionMenu(font_choice, chinese_font_var, *system_fonts).pack()

        # 中文字号输入
        chinese_font_size = askinteger("中文字号设置", "请输入中文字号 (例如: 12):", parent=font_choice)

        # 英文字体选择
        english_font_var = StringVar(font_choice)
        english_font_var.set(system_fonts[0])  # 默认设置为第一个字体
        Label(font_choice, text="选择英文字体:").pack()
        OptionMenu(font_choice, english_font_var, *system_fonts).pack()

        # 英文字号输入
        english_font_size = askinteger("英文字号设置", "请输入英文字号 (例如: 12):", parent=font_choice)

        # 确认按钮
        confirm_button = Button(font_choice, text="确认", command=lambda: callback(
            chinese_font_var.get(), chinese_font_size, english_font_var.get(), english_font_size, font_choice))
        confirm_button.pack(pady=10)

    def format_normal_text_in_document(doc, chinese_font, chinese_font_size, english_font, english_font_size,
                                       font_choice):
        skip = False  # 跳过标志初始化为False
        start_formatting = False  # 开始格式化的标志初始化为False

        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                if "参考文献" in paragraph.text:
                    skip = True  # 开始跳过
                elif "致谢" in paragraph.text:
                    skip = False  # 停止跳过
                    continue  # 确保“致谢”之后的段落不被跳过

            if skip:
                continue  # 如果处于跳过状态，忽略当前段落的处理

        pattern = r'.*引言$'

        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                if re.fullmatch(pattern, paragraph.text.strip()):
                    start_formatting = True  # 开始处理文档

            if start_formatting:
                if paragraph.style.name == 'Normal':
                    if "摘    要" not in paragraph.text and "ABSTRACT" not in paragraph.text:
                        if "目    录" not in paragraph.text:
                            if not has_graphics(paragraph):
                                for run in paragraph.runs:
                                    run.font.name = chinese_font
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                                    run.font.size = Pt(chinese_font_size)
                                    run._element.rPr.rFonts.set(qn('w:ascii'), english_font)
                                    run._element.rPr.rFonts.set(qn('w:hAnsi'), english_font)
                                    run.font.size = Pt(english_font_size)
                                    #
                                    # if paragraph.paragraph_format.first_line_indent is None:
                                    #     paragraph.paragraph_format.first_line_indent = Cm(0.85)
                                    #     paragraph.paragraph_format.left_indent = Cm(0)
                                    # else:
                                    #     if paragraph.paragraph_format.first_line_indent > Cm(0):
                                    #         paragraph.paragraph_format.first_line_indent = Cm(0.85)
                                    #         paragraph.paragraph_format.left_indent = Cm(0)

        font_choice.destroy()  # 关闭字体选择窗口

    def save_document(doc):
        save_path = filedialog.asksaveasfilename(title="保存文档", defaultextension=".docx",
                                                 filetypes=[("Word Documents", "*.docx")])
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("完成", "文件已处理完毕")
        else:
            messagebox.showerror("错误", "未保存文档")

    doc_path = select_document()
    if doc_path:
        doc = Document(doc_path)
        root = tk.Tk()
        root.withdraw()
        get_font_settings(root, lambda cf, cfs, ef, efs, fc: [format_normal_text_in_document(doc, cf, cfs, ef, efs, fc),
                                                              save_document(doc)])
    else:
        messagebox.showerror("错误", "未选择文档，操作取消。")


def apply_custom_template_4():
    def set_normal_style_between_sections(doc):
        in_abstract_to_toc = False  # 控制是否在“摘要”到“目录”之间
        in_main_content = False  # 控制是否在主要内容部分（第一个一级标题之后）
        toc_found = False  # 标记是否找到了“目录”

        for paragraph in doc.paragraphs:
            # 标记“摘要”开始
            if "摘    要" in paragraph.text:
                in_abstract_to_toc = True

            # 标记“目录”找到，并结束“摘要”到“目录”的部分
            if "目    录" in paragraph.text:
                in_abstract_to_toc = False
                toc_found = True

            # 如果找到“目录”且当前段落是一级标题，则开始主要内容部分
            if toc_found and paragraph.style.name == 'Heading 1':
                in_main_content = True
                toc_found = False  # 重置，以防之后有其他内容需要处理

            # 如果在“摘要”到“目录”之间，或者在主要内容部分
            if in_abstract_to_toc or (in_main_content and not paragraph.style.name.startswith('Heading')):
                paragraph.style = doc.styles['Normal']

    def update_headers_if_text_exists(doc, header_text):
        for section in doc.sections:
            if any(paragraph.text.strip() for paragraph in section.header.paragraphs):
                clear_and_set_new_header(section.header, header_text)
            if not section.first_page_header.is_linked_to_previous:
                if any(paragraph.text.strip() for paragraph in section.first_page_header.paragraphs):
                    clear_and_set_new_header(section.first_page_header, header_text)
            if section.even_page_header and not section.even_page_header.is_linked_to_previous:
                if any(paragraph.text.strip() for paragraph in section.even_page_header.paragraphs):
                    clear_and_set_new_header(section.even_page_header, header_text)

    def clear_and_set_new_header(header, text):
        for paragraph in header.paragraphs:
            paragraph.clear()
        new_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        new_paragraph.different_first_page_header_footer = False
        run = new_paragraph.add_run(text)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def format_abstract(doc, text_to_format):
        for paragraph in doc.paragraphs:
            if text_to_format in paragraph.text:
                for run in paragraph.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(16)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_before = Pt(28)
                paragraph.paragraph_format.space_after = Pt(28)

    def format_abstract_in_english(doc, text_to_format):
        for paragraph in doc.paragraphs:
            if text_to_format in paragraph.text:
                for run in paragraph.runs:
                    # 设置字体为Times New Roman
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    # 设置字号为三号字（16磅）
                    run.font.size = Pt(16)
                    # 设置字体加粗
                    run.font.bold = True
                    # 设置居中对齐
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 设置单倍行距
                    paragraph.paragraph_format.line_spacing = 1
                    # 设置段前2行（约28磅）
                    paragraph.paragraph_format.space_before = Pt(28)
                    # 设置段后2行（约28磅）
                    paragraph.paragraph_format.space_after = Pt(28)

    def set_single_line_spacing_for_images(doc):
        """遍历文档中的所有段落，如果包含图片，则设置单倍行距。"""
        for paragraph in doc.paragraphs:
            if has_graphics(paragraph):
                # 设置段落的行距为单倍行距
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def modify_and_format_text_on_first_page(doc, modifications, formats):
        first_page_paragraphs = doc.paragraphs[:7]  # 假设第一页的内容在前10个段落中
        for paragraph in first_page_paragraphs:
            # print("Processing paragraph:", paragraph.text)
            for old_text, new_text in modifications.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
            for text, format_settings in formats.items():
                font, size, bold = format_settings  # 解包格式设置
                if text in paragraph.text:
                    for run in paragraph.runs:
                        if text in run.text:
                            # 设置字体
                            run.font.name = font
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
                            # 设置字号
                            run.font.size = Pt(size)
                            # 设置加粗
                            run.font.bold = bold

    def set_page_break_before(paragraph):
        """为指定段落设置段前分页"""
        p = paragraph._p  # 访问底层的xml元素
        pPr = p.get_or_add_pPr()  # 获取或添加段落属性
        pageBreakBefore = OxmlElement('w:pageBreakBefore')  # 创建段前分页元素
        pPr.append(pageBreakBefore)  # 将段前分页添加到段落属性中

    def format_keywords(doc):
        toc_found = False  # 用于标记是否找到“目    录”
        for paragraph in doc.paragraphs:
            # print("Processing paragraph:", paragraph.text)
            if "目    录" in paragraph.text:
                toc_found = True
                break
            if not has_graphics(paragraph):
                # 处理中文关键词
                if "关键词：" in paragraph.text:
                    process_keywords(paragraph, "关键词：", '黑体', '宋体', False)
                    paragraph.paragraph_format.space_after = Pt(0)
                elif "Key words：" in paragraph.text:
                    process_keywords(paragraph, "Key words：", 'Times New Roman', 'Times New Roman', True)
                    paragraph.paragraph_format.space_after = Pt(0)
            if toc_found:
                break  # 如果已经处理到“目 录”，则停止进一步处理

    def process_keywords(paragraph, keyword_text, keyword_font, text_font, bold):
        # 拆分原始文本以提取和格式化关键词部分
        parts = paragraph.text.split(keyword_text)
        paragraph.clear()

        # 添加关键词前的文本，假设前文使用宋体
        if parts[0].strip():
            add_text_run(paragraph, parts[0], text_font, False)

        # 添加格式化的关键词，使用黑体或Times New Roman
        add_text_run(paragraph, keyword_text, keyword_font, bold)

        # 添加关键词后的文本，使用宋体或Times New Roman
        if len(parts) > 1:
            add_text_run(paragraph, parts[1], text_font, False)

    def add_text_run(paragraph, text, font_name, bold):
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia' if font_name == '黑体' else 'w:ascii'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi' if font_name == 'Times New Roman' else 'w:eastAsia'), font_name)
        run.font.size = Pt(12)  # 设置字号为12磅

    def should_add_page_break(paragraphs, index):
        # 检查前1个段落是否有文本，返回True如果有任何一个包含文本
        start_index = max(index - 3, 0)
        for i in range(start_index, index):
            if paragraphs[i].text.strip():
                return True
        return False

    def add_page_break_before_headings(doc):
        paragraphs = list(doc.paragraphs)
        first_heading_found = False
        for i, paragraph in enumerate(paragraphs):
            if paragraph.style.name == 'Heading 1':
                # print("Processing paragraph:", paragraph.text)
                if not first_heading_found:
                    first_heading_found = True
                    continue  # 跳过文档中的第一个一级标题
                if should_add_page_break(paragraphs, i):
                    set_page_break_before(paragraph)

    def remove_blank_heading_ones(doc):
        # 准备一个列表来收集需要删除的段落的引用
        paragraphs_to_remove = []

        # 遍历所有段落，找到空白的一级标题
        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1' and not paragraph.text.strip():
                # 添加到待删除列表
                paragraphs_to_remove.append(paragraph)

        # 从文档中删除收集的空白段落
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)

    def modify_figure_paragraphs(doc):
        # 正则表达式匹配以“图”开头，后面跟数字的段落
        figure_pattern = r'^(表|图)\d+.*$'

        for paragraph in doc.paragraphs:
            if re.match(figure_pattern, paragraph.text):
                # 对中文和英文字符设置不同字体和字号
                for run in paragraph.runs:
                    if any(char in run.text for char in
                           'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890-'):
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10.5)  # 五号字大约等于10.5磅
                        # 设置固定行距为20磅
                        paragraph.paragraph_format.line_spacing = Pt(20)
                    else:
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)  # 五号字大约等于10.5磅
                        # 设置固定行距为20磅
                        paragraph.paragraph_format.line_spacing = Pt(20)

    def operate_cited(doc):
        start_index, end_index = None, None
        # 一次遍历确定起始和结束索引
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text == "参考文献" and paragraph.style.name == 'Heading 1':
                start_index = i + 1  # 从"参考文献"之后的段落开始处理
            elif text == "致谢" and paragraph.style.name == 'Heading 1':
                end_index = i  # 处理直到"致谢"之前的段落
                break  # 一旦找到两个索引，即停止循环

        for paragraph in doc.paragraphs[start_index:end_index]:
            match = re.search(r'(\D+)(\d{1,2})(\D+)', paragraph.text)
            if match:
                before, number, after = match.groups()
                new_text = before.replace(before[-1], '[') + number + after.replace(after[0], ']')
                paragraph.text = re.sub(r'(\D+)(\d{1,2})(\D+)', new_text, paragraph.text, count=1)
            for run in paragraph.runs:
                if run.text.strip() != '':
                    if any(ord(c) > 128 for c in run.text):  # 判断是否包含中文字符
                        run.font.name = '宋体'
                    else:
                        run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(20)

            indent_xml = """
                           <w:ind xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:left="0" w:hanging="280"/>
                       """
            indent_element = parse_xml(indent_xml)
            paragraph._p.insert(2, indent_element)

            modified_text = paragraph.text.replace('，', ',')
            modified_text = modified_text.replace('。', '.')
            modified_text = modified_text.replace('；', ';')
            modified_text = modified_text.replace('：', ':')
            modified_text = modified_text.replace('【', '[')
            modified_text = modified_text.replace('】', ']')
            modified_text = modified_text.replace('（', '(')
            modified_text = modified_text.replace('）', ')')
            modified_text = modified_text.replace('. ', '.')

            paragraph.text = modified_text

    # def format_specific_keywords_and_text_after_them(doc):
    #     keywords = ["题    目", "学院名称", "专    业", "班    级", "学    号", "学生姓名", "指导教师", "完成日期"]
    #     abstract_found = False  # 标记是否找到了“摘要”
    #
    #     for paragraph in doc.paragraphs:
    #         # 如果找到“摘要”，停止格式化
    #         if "摘    要" in paragraph.text:
    #             abstract_found = True
    #             break
    #
    #         original_text = paragraph.text
    #         for keyword in keywords:
    #             if keyword in original_text:
    #                 # 找到关键字，分割文本
    #                 parts = original_text.split(keyword, 1)
    #                 if len(parts) > 1:
    #                     # 清除原段落内容
    #                     paragraph.clear()
    #                     # 重写并格式化关键字部分（宋体小三号，加粗）
    #                     run = paragraph.add_run(parts[0] + keyword)
    #                     run.font.name = '宋体'
    #                     run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    #                     run.font.size = Pt(16)  # 设置字号为小三号（约16磅）
    #                     run.font.bold = True
    #
    #                     # 重写并格式化关键字之后的部分（楷体小三号）
    #                     run = paragraph.add_run(parts[1])
    #                     run.font.name = 'KaiTi'  # 设置字体为楷体
    #                     run._element.rPr.rFonts.set(qn('w:eastAsia'), 'KaiTi')
    #                     run.font.size = Pt(15)  # 设置字号为小三号（约15磅）
    #                     run.font.bold = False

    def operate_normal_symbol(doc):
        half_to_full_map = {
            ',': '，',
            '?': '？',
            '!': '！',
            ':': '：',
            ';': '；',
            '(': '（',
            ')': '）',
            '{': '｛',
            '}': '｝'
        }

        punctuation_marks = set("。？！，、；：「」『』（）《》【】〈〉“”‘’…—")
        brackets = {'（', '）', '(', ')'}  # 定义括号集合，可根据需要添加全角和半角的其它括号类型

        start_processing = False  # 使用一个标志来控制处理的开始

        pattern = r'.*引言$'

        for paragraph in doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                if re.fullmatch(pattern, paragraph.text.strip()):
                    start_processing = True  # 开始处理文档

            if start_processing and paragraph.style.name == 'Normal':
                if not has_graphics(paragraph):  # 确保不处理包含图形的段落
                    for run in paragraph.runs:
                        # 转换半角到全角符号
                        new_text = ''.join([half_to_full_map.get(char, char) for char in run.text])
                        run.text = new_text

                    # 删除连续的标点符号中的第一个
                    original_text = paragraph.text
                    new_text = []
                    chars = list(original_text)  # 转换为字符列表以处理连续标点
                    i = 0

                    while i < len(chars) - 1:  # 减1避免索引越界
                        # 检查当前和下一个字符是否都在标点集合中，且不含括号
                        if chars[i] in punctuation_marks and chars[i + 1] in punctuation_marks and \
                                (chars[i] not in brackets and chars[i + 1] not in brackets):
                            # 删除当前的重复标点符号，保留下一个（通过不增加i来间接实现）
                            del chars[i]
                        else:
                            i += 1  # 移动到下一个字符

                    paragraph.text = ''.join(chars)  # 更新段落文本

    def add_footer_with_auto_numbering(doc_path):

        word = win32.gencache.EnsureDispatch('Word.Application')
        doc = word.Documents.Open(doc_path)
        word.Visible = False  # 设置为 True 可以看到操作过程，便于调试
        # 查找第一个“标题 1”并插入分节符
        found = False
        for paragraph in doc.Paragraphs:
            if paragraph.Style.NameLocal == '标题 1':  # 根据你的Word版本调整样式名称
                found = True
                # 在“标题 1”起始处插入分节符
                range = paragraph.Range
                range.Collapse(Direction=1)  # Collapse the range to its start
                range.InsertBreak(Type=win32.constants.wdSectionBreakNextPage)
                break
        if found:
            # 获取最新创建的分节
            section = doc.Sections.Last
            # 设置新分节的页脚
            footer = section.Footers(win32.constants.wdHeaderFooterPrimary)
            footer.LinkToPrevious = False  # 不与前一节的页脚链接

            # 重置页码开始为 1
            footer.PageNumbers.RestartNumberingAtSection = True
            footer.PageNumbers.StartingNumber = 1  # 正确设置页码从 1 开始

            # 插入格式化的页码字段，并确保破折号正确插入
            footer_text = footer.Range
            footer_text.Text = " - "  # 首先插入前破折号
            footer_text.Collapse(Direction=1)  # Collapse range to the end
            footer_text.Fields.Add(footer_text, win32.constants.wdFieldEmpty, r'PAGE \* Arabic \* MERGEFORMAT', True)
            footer_text.InsertAfter(" - ")  # 在页码后添加第二个破折号
            footer.Range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
            # 保存更改
            doc.Save()
        # 关闭文档和 Word 应用
        doc.Close()
        word.Quit()

    def set_font_size(paragraph, size):
        for run in paragraph.runs:
            run.font.size = Pt(size)

    def renum(doc_path):
        def is_numbered_paragraph(paragraph):
            # 匹配例如 "（1） ", "（2） ", "（10） " 等格式
            return re.match(r'^（\d+）', paragraph.text)

        # 初始化编号
        current_number = 1

        # 重新编号带有全角括号编号的段落
        for para in doc.paragraphs:
            if is_numbered_paragraph(para):
                # 提取当前编号
                match = re.match(r'^（(\d+)）', para.text)
                if match:
                    original_number = int(match.group(1))
                    if original_number == 1:
                        current_number = 1  # 遇到 "（1）" 则重新开始编号
                    else:
                        # 只对非 "（1）" 的段落重新编号
                        if current_number == 1:
                            current_number += 1
                        para.text = re.sub(r'^（\d+）', f'（{current_number}）', para.text)
                        current_number += 1
                    # 设置段落字体为小四号（12磅）
                    set_font_size(para, 12)

    def modify_table(doc_path):
        table = doc.tables[0]

        # 遍历表格中的行
        for row in table.rows:
            # 获取第二列的单元格
            cell = row.cells[1]
            # 遍历单元格中的所有段落，设置字体和大小
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
                for run in paragraph.runs:
                    run.font.name = '楷体'  # 设置字体为楷体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'KaiTi')
                    run.font.size = Pt(15)  # 设置字体大小为小三号大小

    def set_heading_one_for_specific_paragraphs(doc_path):
        target_texts = ["致谢", "参考文献"]  # 指定需要修改为一级标题的段落文本

        for paragraph in doc.paragraphs:
            if paragraph.text.strip() in target_texts:  # 检查段落文本是否精确匹配列表中的任何一个元素
                # print("Processing paragraph:", paragraph.text)
                paragraph.style = doc.styles['Heading 1']  # 应用一级标题样式
                # 寻找并删除编号相关的XML节点
                pPr = paragraph._element.pPr
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)

    def set_continuous_heading_numbers(doc_path):
        heading_levels = {}  # 用于存储各级标题的当前计数
        unnumbered_headings = ["参考文献", "致谢"]  # 不进行编号的特定标题列表

        for paragraph in doc.paragraphs:
            # 检查段落是否为标题
            if paragraph.style.name.startswith('Heading'):
                # 跳过不编号的特定标题
                if any(unnumbered_heading in paragraph.text for unnumbered_heading in unnumbered_headings):
                    continue

                level = int(paragraph.style.name.split(' ')[1])  # 获取标题级别
                # 初始化或更新标题级别计数
                if level in heading_levels:
                    # 重置下一级标题的计数
                    if level == 1:
                        heading_levels[2] = 0  # 重置二级标题计数
                        heading_levels[3] = 0  # 重置三级标题计数
                    elif level == 2:
                        heading_levels[3] = 0  # 重置三级标题计数
                    heading_levels[level] += 1
                else:
                    # 如果是新的标题级别，重置该级别和所有更低级别的计数
                    for l in range(level, 10):  # 假设标题不会超过9级
                        heading_levels[l] = 0
                    heading_levels[level] = 1

                # 生成新的标题编号
                number = '.'.join(
                    str(heading_levels[l]) for l in sorted(heading_levels) if heading_levels[l] > 0 and l <= level)
                # 更新段落文本以包含新的标题编号
                if paragraph.text.startswith(number):
                    # 如果已有正确的编号，跳过
                    continue
                else:
                    # 移除原有编号
                    text_without_number = paragraph.text.split(' ', 1)[-1] if ' ' in paragraph.text else paragraph.text
                    paragraph.text = f"{number} {text_without_number}"

    def remove_numbering(paragraph):
        """Remove numbering from a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        numPr = pPr.numPr
        if numPr:
            pPr.remove(numPr)

    def add_numbering_between_headings(doc):
        start_numbering = False
        counter = 1  # 编号开始于1
        heading_style = 'Heading 1'  # 一级标题的样式名
        reference_heading_found = False
        pattern = re.compile(r'^\[\d+\]\s*|\d+\.\s*|\(\d+\)\s*')  # 正则表达式识别编号

        for paragraph in doc.paragraphs:
            if '参考文献' in paragraph.text and paragraph.style.name == heading_style:
                start_numbering = True
                reference_heading_found = True
                continue

            if '致谢' in paragraph.text and paragraph.style.name == heading_style and reference_heading_found:
                break

            if start_numbering and reference_heading_found:
                if paragraph.style.name != heading_style and paragraph.text.strip():
                    original_text = pattern.sub('', paragraph.text.strip())
                    remove_numbering(paragraph)  # 移除自动编号
                    paragraph.clear()
                    paragraph.add_run(f"[{counter}] {original_text}")
                    counter += 1
                # print(f"编号 {counter - 1}: {paragraph.text[:30]}")  # 显示编号后的段落前30个字符

                # 设置行间距
                paragraph.paragraph_format.line_spacing = Pt(20)
                paragraph.paragraph_format.first_line_indent = -Cm(0.74)  # 每个字符大约0.1英寸
                paragraph.paragraph_format.left_indent = Cm(0.74)

    def update_paragraphs(doc):
        # 定义用于匹配 "图X-X *" 的正则表达式
        pattern = re.compile(r'^图(\d+)-(\d+)(.*)')

        # 用于记录每个第一个X的最新第二个X值
        figure_count = {}

        # 遍历所有段落
        for para in doc.paragraphs:
            new_text = para.text
            # 查找所有匹配的文本
            matches = pattern.findall(new_text)

            if matches:
                # 初始化临时文本
                temp_text = new_text
                for match in matches:
                    first_x_value = int(match[0])
                    second_x_value = int(match[1])
                    rest_of_text = match[2]

                    # 如果是新的第一个X，初始化计数为1
                    if first_x_value not in figure_count:
                        figure_count[first_x_value] = 1

                    # 获取应当替换的第二个X值
                    expected_x_value = figure_count[first_x_value]

                    # 构建旧的和新的字符串
                    old_str = f'图{first_x_value}-{second_x_value}{rest_of_text}'
                    new_str = f'图{first_x_value}-{expected_x_value}{rest_of_text}'

                    # 只替换第一个匹配的字符串
                    temp_text = temp_text.replace(old_str, new_str, 1)

                    # 更新计数
                    figure_count[first_x_value] += 1

                # 更新段落文本
                para.text = temp_text

    def format_normal_text_in_abstract(doc, start_paragraph_index):
        skip = False
        for paragraph in doc.paragraphs[start_paragraph_index:]:
            if "摘    要" in paragraph.text:
                skip = False
            elif "目    录" in paragraph.text:
                skip = True

            if skip:
                break

            if paragraph.style.name == 'Normal':
                if paragraph.style.name == 'Normal' and "摘    要" not in paragraph.text and "ABSTRACT" not in paragraph.text:
                    if paragraph.style.name == 'Normal' and "目    录" not in paragraph.text:
                        # print("Processing paragraph:", paragraph.text)
                        for run in paragraph.runs:
                            if not run.font.superscript:
                                # 设置中文字符为宋体小四号字
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                # 设置英文字符为Times New Roman小四号字
                                run.font.name = 'Times New Roman'
                                run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                                run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                                # 设置字号为小四号字（12磅）
                                run.font.size = Pt(12)
                                # 设置行距为固定值20磅
                                paragraph.paragraph_format.line_spacing = Pt(20)
                                if paragraph.paragraph_format.first_line_indent is None:
                                    paragraph.paragraph_format.first_line_indent = Cm(0.85)
                                    paragraph.paragraph_format.left_indent = Cm(0)
                                else:
                                    if paragraph.paragraph_format.first_line_indent > Cm(0):
                                        paragraph.paragraph_format.first_line_indent = Cm(0.85)
                                        paragraph.paragraph_format.left_indent = Cm(0)

    # print("Processing paragraph:", paragraph.text)
    selected_doc_path = select_document()
    if selected_doc_path:
        doc = Document(selected_doc_path)
        messagebox.showinfo('Info', "处理文档需要一些时间，请耐心等待")

        # 格式化表格
        modify_table(selected_doc_path)

        # 删除所有空白的一级标题
        remove_blank_heading_ones(doc)

        # 格式化文献格式
        operate_cited(doc)

        set_heading_one_for_specific_paragraphs(selected_doc_path)
        set_continuous_heading_numbers(selected_doc_path)
        # add_numbering_between_headings(selected_doc_path)
        add_numbering_between_headings(doc)
        update_paragraphs(doc)

        # 设置文档段落为Normal样式
        set_normal_style_between_sections(doc)

        header_text_one = "杭州电子科技大学信息工程学院本科毕业设计"
        # 格式化页眉
        header_text = header_text_one
        update_headers_if_text_exists(doc, header_text)

        format_normal_text_in_abstract(doc, 5)
        # 需要修改的文本，字体和字号
        modifications = {
            "杭州电子科技大学信息工程学院": "杭州电子科技大学信息工程学院",
            "本科毕业设计": "本科毕业设计",
            "（2024届）": "（2024届）"
        }
        # 要设置的文本，字体和字号
        formats = {
            "杭州电子科技大学信息工程学院": ("宋体", 28, True),
            "本科毕业设计": ("宋体", 28, True),
            "（2024届）": ("宋体", 22, True)  # 二号字号一般为22磅
        }

        # 修改并格式化第一页中的特定文本
        modify_and_format_text_on_first_page(doc, modifications, formats)

        # 摘要修改
        text_to_format1 = "摘    要"
        format_abstract(doc, text_to_format1)
        text_to_format2 = "ABSTRACT"
        format_abstract_in_english(doc, text_to_format2)
        text_to_format3 = "目    录"
        format_abstract(doc, text_to_format3)

        # 操作标点
        operate_normal_symbol(doc)

        # 修改图注
        modify_figure_paragraphs(doc)

        # 图片设置为单倍行距
        set_single_line_spacing_for_images(doc)

        # 格式化关键词
        format_keywords(doc)

        # 为符合条件的一级标题添加段前分页
        add_page_break_before_headings(doc)

        renum(doc)

        # 另存为文件
        new_doc_path = select_save_as()
        if new_doc_path:
            doc.save(new_doc_path)
            add_footer_with_auto_numbering(new_doc_path)
            messagebox.showinfo("Info", "文件已处理完毕")
        else:
            messagebox.showinfo("Info", "取消保存文件")
    else:
        messagebox.showinfo("Info", "未选择文件或者取消")


def open_custom_template_window(root):
    # 隐藏主窗口
    root.withdraw()

    # 创建新窗口
    new_window = tk.Toplevel(root)
    new_window.title("自定义模板处理")
    new_window.geometry("350x350")

    btn_step1 = tk.Button(new_window, text="设置各级标题样式", command=apply_custom_template_1, width=20, height=2)
    btn_step1.pack(pady=10)

    btn_step2 = tk.Button(new_window, text="设置页面布局", command=apply_custom_template_2, width=20, height=2)
    btn_step2.pack(pady=10)

    btn_step3 = tk.Button(new_window, text="设置正文中英文字体字号", command=apply_custom_template_3, width=20,
                          height=2)
    btn_step3.pack(pady=10)

    btn_step4 = tk.Button(new_window, text="校勘文档中的错误", command=apply_custom_template_4, width=20, height=2)
    btn_step4.pack(pady=10)

    btn_return = tk.Button(new_window, text="返回主菜单", command=lambda: close_custom_window(root, new_window),
                           width=20, height=2)
    btn_return.pack(pady=10)


def create_ui():
    root = tk.Tk()
    root.title("毕设论文格式校勘")
    root.geometry("300x100")

    messagebox.showinfo("Info", "本工具默认模板为：理工类专业毕业设计模板-240327修订V1.2.doc")
    btn_default_template = tk.Button(root, text="使用默认模板处理文档", command=apply_default_template)
    btn_default_template.pack(pady=10)

    btn_open_custom_templates = tk.Button(root, text="打开自定义模板处理窗口",
                                          command=lambda: open_custom_template_window(root))
    btn_open_custom_templates.pack(pady=10)

    root.mainloop()


def close_custom_window(root, new_window):
    # 销毁新窗口
    new_window.destroy()
    # 重新显示主窗口
    root.deiconify()


if __name__ == "__main__":
    clear_com_cache()
    create_ui()
